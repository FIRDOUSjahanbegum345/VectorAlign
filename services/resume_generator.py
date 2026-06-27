import os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

# =========================
# CONFIG
# =========================
GENERATED_DIR = "generated"
os.makedirs(GENERATED_DIR, exist_ok=True)

# =========================
# TEXT CLEANER
# =========================
def clean_resume_text(text):
    """
    Formats AI-generated resume text for readability
    """
    if not text:
        return ""

    text = text.replace("\r", "")
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    formatted_lines = []

    for line in lines:
        if any(keyword in line.lower() for keyword in [
            "skills", "projects", "education", "experience",
            "certifications", "summary", "work"
        ]):
            formatted_lines.append("\n" + line.upper())
        else:
            formatted_lines.append(line)

    return "\n".join(formatted_lines)

# =========================
# PDF BUILDER
# =========================
def build_pdf(pdf_path, optimized_resume, feedback):
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph("AI Resume Optimization Report", styles['Title']))
    story.append(Spacer(1, 12))

    # ATS Score Analysis
    story.append(Paragraph("ATS Score Analysis", styles['Heading2']))
    story.append(Paragraph(feedback.get("ats_score_evaluation", "N/A"), styles['Normal']))
    story.append(Spacer(1, 12))

    # Optimized Resume
    story.append(Paragraph("Professional Optimized Resume", styles['Heading2']))
    story.append(Paragraph(optimized_resume, styles['Normal']))
    story.append(Spacer(1, 12))

    # Missing Keywords
    story.append(Paragraph("Missing Keywords", styles['Heading2']))
    missing = feedback.get("missing_keywords", [])
    if missing:
        items = [ListItem(Paragraph(f"{k}", styles['Normal'])) for k in missing]
        story.append(ListFlowable(items, bulletType='bullet'))
    else:
        story.append(Paragraph("No missing keywords detected.", styles['Normal']))
    story.append(Spacer(1, 12))

    # Structural Feedback
    story.append(Paragraph("Structural Feedback", styles['Heading2']))
    structural = feedback.get("structural_feedback", [])
    if structural:
        items = [ListItem(Paragraph(f"{i}", styles['Normal'])) for i in structural]
        story.append(ListFlowable(items, bulletType='bullet'))
    else:
        story.append(Paragraph("Resume structure is strong.", styles['Normal']))
    story.append(Spacer(1, 12))

    # Impact Suggestions
    story.append(Paragraph("Impact Improvements", styles['Heading2']))
    impact = feedback.get("quantifiable_impact_suggestions", [])
    if impact:
        items = [ListItem(Paragraph(f"{i}", styles['Normal'])) for i in impact]
        story.append(ListFlowable(items, bulletType='bullet'))
    else:
        story.append(Paragraph("No improvements required.", styles['Normal']))

    doc.build(story)

# =========================
# MAIN GENERATOR
# =========================
def create_optimized_resume(optimized_content, feedback, file_id):
    # ---------------- GET RESUME ----------------
    optimized_resume = feedback.get("optimized_resume", "")
    if not optimized_resume:
        optimized_resume = feedback.get("tailored_summary_suggestion", "")

    optimized_resume = clean_resume_text(optimized_resume)

    # ---------------- DOCX ----------------
    doc = Document()
    doc.add_heading("AI Resume Optimization Report", 0)

    doc.add_heading("ATS Score Analysis", level=1)
    doc.add_paragraph(feedback.get("ats_score_evaluation", "N/A"))

    doc.add_heading("Professional Optimized Resume", level=1)
    doc.add_paragraph(optimized_resume)

    # Missing Keywords
    doc.add_heading("Missing Keywords", level=1)
    missing = feedback.get("missing_keywords", [])
    if missing:
        for k in missing:
            doc.add_paragraph(f"• {k}")
    else:
        doc.add_paragraph("No missing keywords detected.")

    # Structural Feedback
    doc.add_heading("Structural Feedback", level=1)
    structural = feedback.get("structural_feedback", [])
    if structural:
        for i in structural:
            doc.add_paragraph(f"• {i}")
    else:
        doc.add_paragraph("Resume structure is strong.")

    # Impact Suggestions
    doc.add_heading("Impact Improvements", level=1)
    impact = feedback.get("quantifiable_impact_suggestions", [])
    if impact:
        for i in impact:
            doc.add_paragraph(f"• {i}")
    else:
        doc.add_paragraph("No improvements required.")

    docx_path = os.path.join(GENERATED_DIR, f"optimized_{file_id}.docx")
    pdf_path = os.path.join(GENERATED_DIR, f"optimized_{file_id}.pdf")

    doc.save(docx_path)

    # ---------------- PDF (structured with ReportLab) ----------------
    try:
        build_pdf(pdf_path, optimized_resume, feedback)
    except Exception as e:
        with open(pdf_path, "w", encoding="utf-8") as f:
            f.write(optimized_resume)

    # ---------------- RESPONSE ----------------
    return {
        "docx_path": f"/generated/optimized_{file_id}.docx",
        "pdf_path": f"/generated/optimized_{file_id}.pdf",
        "optimized_resume": optimized_resume,
        "keyword_match_score": feedback.get("keyword_match_score", 0),
        "job_match_score": feedback.get("job_match_score", 0)
    }
